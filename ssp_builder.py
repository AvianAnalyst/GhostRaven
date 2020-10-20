# TODO: factor out yaml interpreter class
# TODO: write tests
# TODO: convert yaml to lists
# todo: write argparse to run from commandline with flags
import re
import docx
from collections import defaultdict
from logging import warning



class SSPBuilder:

    def __init__(self, env, contents: dict):
        if env == 'dod':
            self.template = docx.Document('dod_template.docx')
        else:
            self.template = docx.Document('template.docx')
        self.contents = contents

        self.undefined_keys = defaultdict(lambda: 0)

    def _replace_text(self, match, content):
        key = match.group(0).strip('{}')
        try:
            return content[key]
        except KeyError:
            self.undefined_keys[key] += 1
            warning(key)
            return '' # hides mistakes made

    def build_ssp(self):
        for table in self.template.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell.text = re.sub(r'{{(.+?)}}', lambda match: self._replace_text(match, self.contents), cell.text)

    def save(self, output_name = 'output/working_output.docx'):
        self.template.save(output_name)
        if self.undefined_keys:
            with open('output/error_log.txt', 'w') as f:
                log = 'The following keys were referenced in the template, or yaml file, but were not defined in any yaml file:\n'
                log += '\n'.join([str(key).strip('()') for key in zip(self.undefined_keys.keys(), self.undefined_keys.values())])
                f.write(log)
            warning("There were some keys referenced that weren't defined. Please view output/error_log.txt for more information.")




