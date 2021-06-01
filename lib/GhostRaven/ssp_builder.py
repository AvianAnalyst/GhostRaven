# TODO: factor out yaml interpreter class
# TODO: write tests
# TODO: convert yaml to lists
# todo: write argparse to run from commandline with flags
import re
import docx
from collections import defaultdict
from logging import warning

from docx.shared import RGBColor


class SSPBuilder:

    def __init__(self, env_template: str, contents: dict):
        self.template = docx.Document(env_template)

        self.contents = contents

        self.undefined_keys = defaultdict(lambda: 0)

    def _replace_text(self, match, content):
        key = match.group(0).strip('{}')
        try:
            return content[key]
        except KeyError:
            self.undefined_keys[key] += 1
            warning(key)
            return '' # hides mistakes made

    def build_cell(self, cell):
        cell.text = cell.text.lstrip('\n')  # Makes the entire cell one run, which they should be at this point,
                                            # but can be broken by editing word
        try:  # try block to bypass cells that don't request a value
            if re.match(r'^[A-Z]{2}-[0-9]{1,2}(?: \([0-9]{1,2}\))? What is the solution and how is it implemented\?$', cell.text) or \
               re.match(r'^[A-Z]{2}-[0-9]{1,2}(?: \([0-9]{1,2}\))?', cell.text) or \
               re.match(r'Control Summary Information', cell.text):
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
            if re.match(r'^Part [a-z]$', cell.text):  # ensures part columns will be bolded
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
            elif re.search(r'^.+:\n',
                           text := self.contents[re.search(r'{{(.+?)}}', cell.text).group(1)],
                           re.MULTILINE):  # looks for any line that ends in a colon and new line in the content
                working_p = cell.paragraphs[0].clear()  # clears cell contents to start over
                sections = re.split(r'(^.+:\n)', text, flags=re.MULTILINE)  # creates a list of the 'sections' of headings and their content
                # Builds out the sections bolding the headings
                for text in sections:
                    if len(text) > 1 and text[-2] == ':':
                        working_p.add_run(text).bold = True
                    else:
                        working_p.add_run(text)
            else:  # if there is no need to bold anything we can do a simple text replace
                cell.text = re.sub(r'{{(.+?)}}', lambda match: self._replace_text(match, self.contents), cell.text)
        except AttributeError:
            pass  # we don't want to do anything this will trigger on things like table headings

    def build_ssp(self):
        for table in self.template.tables:
            for row in table.rows:
                for cell in row.cells:
                    self.build_cell(cell)

    def save(self, output_name = 'output/working_output.docx'):
        self.template.save(output_name)
        if self.undefined_keys:
            with open('output/error_log.txt', 'w') as f:
                log = 'The following keys were referenced in the template, or yaml file, but were not defined in any yaml file:\n'
                log += '\n'.join([str(key).strip('()') for key in zip(self.undefined_keys.keys(), self.undefined_keys.values())])
                f.write(log)
            warning("There were some keys referenced that weren't defined. Please view output/error_log.txt for more information.")




